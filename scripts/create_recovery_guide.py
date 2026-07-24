from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Neerus-Kitchen-Technical-Recovery-Guide.docx"
TODAY = date.today().isoformat()

INK = "17211B"
GREEN = "237149"
ORANGE = "E95B2A"
MUTED = "66736A"
PALE_GREEN = "EAF7EF"
PALE_ORANGE = "FFF0E9"
PALE_GRAY = "F2F5F3"
LINE = "D7E0DA"
WIDTH = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_run(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def write_cell(cell, text, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
    r = p.add_run(text)
    set_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run(r, size=10.5)


def add_numbered(doc, items):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [500, 8860])
    for index, item in enumerate(items, start=1):
        row = table.add_row()
        prevent_row_split(row)
        set_cell_shading(row.cells[0], PALE_GRAY)
        write_cell(row.cells[0], str(index), bold=True, color=GREEN, size=10)
        write_cell(row.cells[1], item, size=10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=PALE_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, value in zip(header.cells, headers):
        set_cell_shading(cell, header_fill)
        write_cell(cell, value, bold=True, color=GREEN, size=9)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for cell, value in zip(row.cells, row_values):
            write_cell(cell, value)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_callout(doc, title, body, fill=PALE_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, size=10.5, color=GREEN if fill == PALE_GREEN else ORANGE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2, size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_diagram(doc, title, nodes, fills):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_run(r, size=9, color=MUTED, bold=True)
    table = doc.add_table(rows=1, cols=len(nodes))
    set_table_geometry(table, [WIDTH // len(nodes)] * len(nodes))
    for index, (cell, text) in enumerate(zip(table.rows[0].cells, nodes)):
        set_cell_shading(cell, fills[index])
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run(r, size=9, color=INK, bold=True)
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)
    r = caption.add_run(f"Diagram width: 6.5 in (9360 DXA); {len(nodes)} equal columns.")
    set_run(r, size=8.5, color=MUTED, italic=True)


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in ((1, 16, GREEN, 16, 8), (2, 13, GREEN, 12, 6), (3, 11, ORANGE, 9, 5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("NEERU’S HOME KITCHEN  |  TECHNICAL RECOVERY GUIDE")
    set_run(r, size=8, color=MUTED, bold=True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Internal recovery reference · Snapshot prepared {TODAY} · Keep customer backups private")
    set_run(r, size=8, color=MUTED)


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)

    # Editorial cover pattern, paired with the compact_reference_guide preset.
    doc.add_paragraph().paragraph_format.space_after = Pt(44)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NEERU’S HOME KITCHEN")
    set_run(r, size=12, color=ORANGE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Technical Recovery\n& Operations Guide")
    set_run(r, size=28, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("A rebuild-ready reference for the family order desk and customer storefront")
    set_run(r, size=12, color=MUTED)
    add_callout(doc, "Snapshot purpose", "This guide records how the working app is structured, operated, backed up and rebuilt. It deliberately excludes passwords, Supabase service-role keys, VAPID private keys, customer data and other secrets. The Git repository plus a private portable backup are the paired recovery source of truth.", PALE_ORANGE)
    add_table(doc, ["Reference", "Current value"], [
        ("Prepared", TODAY),
        ("Repository", "github.com/neofyne/NeerusKitchen"),
        ("Primary deployment", "Cloudflare Pages: neerushomekitchen.pages.dev"),
        ("Application surfaces", "Customer storefront / and family admin desk /admin"),
        ("Document preset", "compact_reference_guide with editorial_cover opening"),
    ], [2500, 6860])
    doc.add_page_break()

    add_heading(doc, "1. What this app does")
    add_body(doc, "Neeru’s Home Kitchen is a two-surface ordering system for a small residential community. The customer storefront lets residents browse the live menu, search and filter dishes, place an order and pay by UPI. The family order desk lets the kitchen create and manage direct orders, track delivery/payment state, share payment notes, manage dishes/categories, edit customer information and operate the storefront.")
    add_callout(doc, "Operating principle", "Orders are intentionally itemized end-to-end. A person seeing an order card, payment link or WhatsApp payment note should be able to understand what was ordered, how much each item costs and the total without needing an internal code.")

    add_heading(doc, "2. System architecture")
    add_diagram(doc, "Architecture overview", ["Customer storefront\n/", "Admin desk\n/admin", "Cloudflare Pages\nstatic app + edge API", "Supabase\nauth + data"], [PALE_GREEN, PALE_ORANGE, PALE_GRAY, PALE_GREEN])
    add_body(doc, "The React application is compiled by Vite into two HTML entry points. Cloudflare Pages serves the current static build, applies no-cache headers to the HTML shell, and retains long-lived cache headers only for hashed assets and static food images. Supabase is the transactional source for orders, menu state, customer records, settings and push subscriptions.")
    add_diagram(doc, "Media and alert paths", ["Admin image upload", "Cloudflare photo API", "R2 for new banners\nlegacy API fallback for older media", "Customer/admin display"], [PALE_ORANGE, PALE_GRAY, PALE_GREEN, PALE_GREEN])
    add_body(doc, "New announcement banners are stored through the Cloudflare photo route in R2. Existing menu, order and payment image handling is deliberately forwarded to the former API for compatibility. The UI must always handle a missing image gracefully: preview errors return to the upload state and storefront banner errors hide the image instead of showing a broken placeholder.")

    add_heading(doc, "3. Stack and runtime")
    add_table(doc, ["Layer", "Technology", "Responsibility"], [
        ("Client", "React + TypeScript", "Admin desk and customer storefront UI; separate authenticated Supabase clients."),
        ("Build", "Vite", "Compiles index.html and admin.html with hashed JS/CSS assets."),
        ("Database & auth", "Supabase / PostgreSQL / Auth / RLS", "Orders, menu, profiles, settings, admin controls and secure RPCs."),
        ("Hosting", "Cloudflare Pages", "Production static hosting, headers, routes and Pages Functions."),
        ("Object storage", "Cloudflare R2 + legacy media API", "Banner storage in R2; old media remains reachable through compatibility proxy."),
        ("Notifications", "Web Push + Cloudflare Worker", "Private scheduled delivery alerts to enabled family devices."),
        ("Payment", "UPI deep links, QR code, named app choices", "No card processing is held by the app; chosen UPI app handles payment."),
    ], [1700, 2300, 5360])
    add_heading(doc, "4. Routes and page logic")
    add_table(doc, ["Route / surface", "Audience", "Core logic"], [
        ("/", "Customers", "Loads current menu/settings, optional banner/message/headline, categories, search, cart and checkout."),
        ("/admin", "Family admins", "Requires admin session; daily operations, direct-order creation, delivery/payment management, client directory, menu and settings."),
        ("/c/:category/:hero?", "Shared customers", "A category/share link that presents the related items and order path."),
        ("/p/:shareCode", "Customer receiving payment note", "Private payment-note page generated from selected outstanding orders; presents UPI choices and the current payment total."),
        ("/pay/:orderCode/:amount", "Compatibility payment link", "Legacy/simple payment entry route maintained for existing shares."),
    ], [2200, 1700, 5460])
    add_heading(doc, "5. Customer storefront behaviour")
    add_bullets(doc, [
        "Storefront settings control whether orders, banner image, customer message, announcement headline and announcement details are visible. Each content block is optional and can be hidden independently.",
        "Menu search returns a single full-width card so a result and its image are easy to read. Category and full-menu browsing use two columns on phone-sized layouts; two-column image areas are deliberately taller to avoid excessive crop.",
        "The customer session persists on the same device under neerus-kitchen-customer-auth. A customer phone number is the identity safeguard for guest orders; matching phone records prevent duplicate clients despite a spelling change.",
        "Checkout accepts delivery time and a per-order note as optional details. It uses an accessible time picker with 12-hour display in the app.",
        "Payment screens show QR and UPI app choices. The old UPI-reference confirmation field is intentionally hidden in the current customer flow; customers can confirm payment with Neeru directly.",
    ])

    add_heading(doc, "6. Family admin desk behaviour")
    add_table(doc, ["Area", "Working logic"], [
        ("Orders", "One chronological operational board with All orders, To deliver, Delivered and Pending filters. Date controls change the actual date/range shown; card labels are state-aware."),
        ("New/edit order", "Create a direct order without requiring a remark, reminder or time. Select dishes by category/search; add to unpaid orders, or create a distinct new order if the prior one has been paid."),
        ("Order cards", "Show customer, flat, payment/delivery status, readable item lines and item calculations. Card navigation is limited to deliberate controls; editing is not triggered by the entire card."),
        ("Payments", "Pending views and customer history can select unpaid orders and generate a personal shareable payment note. It itemizes each dish, quantity × rate = amount, then the combined total and payment page URL."),
        ("Client history", "Open from an order/customer. History offers compact Today, Yesterday, 7, 15, 30 days and Custom ranges, then supports sharing the selected unpaid orders."),
        ("Client directory", "Search by name/flat/tower; edit existing client details or add a client before their first order. Sort by spend, recent, name, pending or tower/flat, with A/B/C/D refinement."),
        ("Menu", "Edit/search dishes, images, prices, availability, featured state and categories. Full menu/category views use two columns; a typed search uses a full-width single result."),
        ("Settings", "Storefront controls, UPI/QR, WhatsApp contact, push alerts, data export/restore and refresh/update. Refresh is one simple action with a small information popout."),
    ], [2050, 7310])

    add_heading(doc, "7. Core workflows")
    add_diagram(doc, "Customer order lifecycle", ["Browse/search", "Cart", "Checkout", "Guest/auth order RPC", "Admin receives order"], [PALE_GREEN, PALE_GRAY, PALE_ORANGE, PALE_GREEN, PALE_GREEN])
    add_numbered(doc, [
        "A customer browses the daily menu or a shared category/dish link, then chooses quantities.",
        "The checkout captures name, tower, flat and phone; delivery time/note are optional.",
        "The database RPC writes the parent order plus order_items in one controlled operation and uses phone identity to update/claim the customer profile safely.",
        "The admin desk refreshes its order list and may make an in-app alert sound after interaction has unlocked audio. Push alerts are a separate opt-in device capability.",
        "When payment is outstanding, an admin manually uses Share reminder/payment note; the app invokes the platform share sheet rather than sending messages automatically.",
    ])
    add_diagram(doc, "Direct-order and payment-note workflow", ["Admin creates / adds dishes", "Outstanding order(s)", "Select by pending or history range", "Generate /p share code", "Manual share sheet"], [PALE_ORANGE, PALE_GRAY, PALE_GREEN, PALE_GREEN, PALE_ORANGE])
    add_body(doc, "The share text has WhatsApp formatting only when WhatsApp is the target. The app does not auto-send WhatsApp, Instagram or iMessage messages. For safety and portability it uses a standard native share payload, with a plain-text fallback for non-WhatsApp targets.")

    add_heading(doc, "8. Data model and source of truth")
    add_table(doc, ["Data group", "Main records", "Notes"], [
        ("Orders", "orders, order_items", "One order header with an itemized immutable-at-order-time line snapshot: name, unit price, quantity and unit label."),
        ("Menu", "menu_items, daily_menu, dish_categories, menu_item_categories", "Catalog information is separate from the daily availability/price/featured state."),
        ("Customers", "customer_profiles, guest_customer_contacts, customer_admin_notes, restored_customer_profiles", "Profiles connect to authenticated users when possible; guest phone identity supports a small, low-friction community rollout."),
        ("Storefront", "storefront_settings", "Ordering status, UPI/payee, WhatsApp contact and optional announcement/display fields."),
        ("Payments", "payment_note_requests", "Share code and selected outstanding order set for a private payment-note page."),
        ("Alerts", "admin_push_subscriptions, delivery_reminder_push_runs", "Subscribed family devices and deduped delivery-reminder processing."),
        ("Access", "admin_users plus Supabase auth.users", "Admin authority is checked in database RPC/RLS rather than trusted from the client UI."),
    ], [1900, 3050, 4410])
    add_callout(doc, "Data rule", "Supabase is the live structured-data source. Git stores code and migrations, never the live customer/order data. A portable backup is private and contains personal details, addresses, order history and payment references.")
    add_heading(doc, "9. Database migrations and RPCs")
    add_body(doc, "Run the generated new-project setup SQL once for a clean Supabase project. It is assembled at build time by scripts/build-supabase-setup.mjs from the individual migration files. Keep the migration files in version control; they are the code-level database history.")
    add_table(doc, ["Migration family", "What it establishes"], [
        ("schema.sql, customer_storefront.sql, storefront_hardening.sql", "Core orders/menu schema, RLS and safe public storefront query/order placement."),
        ("dish_categories_and_units.sql, multi_dish_categories.sql, dish_promotions.sql", "Category architecture, portions/units and customer sharing/promotion controls."),
        ("manual_order_phone.sql, guest_customer_checkout.sql, instant_customer_access.sql", "Direct/admin and low-friction customer order identity flows."),
        ("two_state_orders.sql, delivery_reminders.sql, push_reminders.sql", "Delivery state, reminder times and opt-in device push lifecycle."),
        ("customer_directory.sql, customer_admin_notes.sql", "Directory metrics/search/update RPCs and private admin notes."),
        ("portable_backup_restore.sql", "Administrator-only JSON backup and transactional restore/reconnection of customer profiles."),
        ("storefront_announcements.sql, storefront_display_controls.sql", "Optional storefront banner/message/title/detail settings."),
    ], [3700, 5660])
    add_body(doc, "Important RPCs include kitchen_today, is_admin, place_customer_order, place_guest_customer_order, create_payment_note_request, get_payment_note_request, admin_customer_directory, admin_update_customer_directory_entry, create_portable_backup, restore_portable_backup and claim_due_delivery_reminders. Exact signatures and grants are preserved in public/supabase-new-project-setup.sql and the source migrations.")

    add_heading(doc, "10. Media, brand and visual dimensions")
    add_table(doc, ["Asset / layout", "Specification", "Reason"], [
        ("Brand logo", "public/neeru-logo.png", "Single current logo asset used by the customer and admin shells."),
        ("Banner upload", "JPG/PNG/WebP/HEIC/HEIF; max 512 KB", "Stored as an R2 key under banner/{admin-id}/{uuid}. Missing image must fall back cleanly."),
        ("Announcement preview", "Minimum 150 px high; banner display max-height 160 px in Settings", "Keeps the manager easy to scan without letting a poster consume the whole settings page."),
        ("Admin dish picker", "Two-column cards; 96 px image row", "Shows more food photo when selecting dishes for an order."),
        ("Customer two-column card", "Near-square image area: 1.06:1 at phone layout; category override 1.22:1 desktop", "Preserves food framing while keeping two cards readable."),
        ("Search result card", "Single column; image uses natural height / contain", "Avoids crop for a single manually searched result."),
        ("Architecture diagrams in this guide", "6.5 in × approximately 0.45–0.8 in, fixed 9360 DXA width", "Printable recovery reference diagrams, not application UI assets."),
    ], [2700, 3500, 3160])

    add_heading(doc, "11. Authentication, privacy and safety")
    add_bullets(doc, [
        "Admin and customer sessions use separate local-storage keys, so an admin can test the storefront without replacing the family session.",
        "RLS and RPC checks are the actual authorization boundary. Never move service-role secrets or VAPID private keys into Vite client variables or Git.",
        "Customer phone data is operational contact data. Treat portable backups, CSV/JSON exports and screenshots as private records.",
        "The app uses manual sharing for customer messaging; it does not silently send payment reminders. This prevents accidental spam and leaves the sender in control.",
        "The user-facing payment reference field is currently hidden. Payment confirmation is handled by direct customer communication until a verified end-to-end process is intentionally reintroduced.",
    ])

    add_heading(doc, "12. Deploy, verify and update")
    add_numbered(doc, [
        "Install dependencies with npm ci (or npm install for a fresh editable checkout). Configure only public VITE_* values locally; use Cloudflare/worker secrets for private keys.",
        "Run npm run dev -- --host 127.0.0.1 and check both / and /admin. Local Vite currently proxies /api requests to the legacy compatibility API; production uses the deployed Cloudflare route.",
        "Run npm run build. The build first creates public/supabase-new-project-setup.sql, then type-checks and bundles both app entries.",
        "Run git diff --check. Confirm no secrets, private backup JSON or customer export files are staged.",
        "Deploy the current dist output to Cloudflare Pages only after explicit approval. Verify fresh HTML headers with Cache-Control: no-cache and verify current hashed assets, not just a cached mobile tab.",
        "For an old/stale device, use the Settings refresh action once; the no-store HTML and service worker update strategy should bring the current build after reload.",
    ])
    add_callout(doc, "Production verification checklist", "Check customer menu load, banner fallback, search/category card layouts, guest checkout, admin sign-in, new direct order with blank optional fields, payment-note share, client directory, image upload, and settings refresh. Confirm the deployed Cloudflare URL—not the older Netlify domain—is the customer-facing link.", PALE_ORANGE)

    add_heading(doc, "13. Disaster recovery")
    add_numbered(doc, [
        "Keep a private portable backup from Settings → Full backup & restore before major database/project changes. Store it off-device in a private location.",
        "Recover the Git repository first. It contains the app, migrations, worker and deployment settings needed to rebuild the codebase.",
        "Create a new Supabase project. In Supabase Authentication, create the intended kitchen administrator, then run public/supabase-new-project-setup.sql in the SQL editor.",
        "Set the new VITE_SUPABASE_URL and VITE_SUPABASE_PUBLISHABLE_KEY in the deployment environment and build/deploy the app.",
        "Sign into /admin, use Settings → Full backup & restore, select the private JSON and restore after reviewing the counts. The restore is transactional: it rolls back entirely if an error occurs.",
        "Reconnect customer accounts when they register using the matching email/phone. Authentication passwords cannot be exported; customer profile and order data can be reclaimed safely.",
        "If moving to a different site/storage account, separately migrate old media. Portable backup stores opaque media references but does not copy underlying legacy Blob/R2 bytes.",
    ])
    add_heading(doc, "14. Secrets and configuration inventory")
    add_table(doc, ["Variable / binding", "Where it belongs", "Do not commit"], [
        ("VITE_SUPABASE_URL", "Local .env.local and Cloudflare Pages build environment", "Public project URL is not secret, but keep configuration per environment."),
        ("VITE_SUPABASE_PUBLISHABLE_KEY", "Local .env.local and Cloudflare Pages build environment", "Publishable client key only; never substitute service-role key."),
        ("PUSH_VAPID_PUBLIC_KEY", "Cloudflare Pages environment or push config API", "Safe to expose to subscribed browser clients."),
        ("VAPID_PRIVATE_KEY, SUPABASE_SERVICE_ROLE_KEY", "Cloudflare Worker secrets only", "Never expose in Vite bundle, repository or document."),
        ("BANNER_PHOTOS", "Cloudflare R2 binding", "Production object store binding; ensure it exists before enabling new banner upload."),
        ("Optional Meta WhatsApp values", "Server-side environment only", "Not required for manual sharing; never put access token in browser app."),
    ], [2900, 4000, 2460])

    add_heading(doc, "15. Repository map")
    add_table(doc, ["Path", "Why it matters"], [
        ("src/main.tsx", "Admin app, direct-order workflow, payment notes, directory and settings."),
        ("src/storefront.tsx", "Customer storefront, browse/search/cart/checkout/payment pages."),
        ("src/styles.css and src/storefront.css", "Responsive visual system, card layouts and mobile behaviour."),
        ("src/supabase.ts", "Separate admin/customer Supabase session clients."),
        ("src/pushNotifications.ts and public/admin-sw.js", "Browser push subscription and notification display."),
        ("functions/api/photos.ts", "Cloudflare photo endpoint; R2 banner handling plus legacy media compatibility."),
        ("workers/reminder-push.ts", "Scheduled web-push dispatch for due delivery reminders."),
        ("supabase/*.sql", "Incremental data/schema/RPC/RLS migrations."),
        ("public/supabase-new-project-setup.sql", "Generated clean-project bootstrap script; regenerate during npm run build."),
        ("public/_headers and public/site.webmanifest", "Cache control and installable-app metadata."),
        ("docs/Neerus-Kitchen-Technical-Recovery-Guide.docx", "This recovery guide; update after material architecture changes."),
    ], [3250, 6110])
    add_heading(doc, "16. Maintenance cadence")
    add_bullets(doc, [
        "Before a major change: create a portable backup, check current deployment and use a local branch/worktree with npm run build.",
        "After a user-facing change: test both customer and admin on a phone-sized layout, especially menu-card crop, popup height, sharing text and optional fields.",
        "After any database migration: include the source SQL, rebuild the generated new-project setup SQL, and run the migration in Supabase before declaring the UI feature live.",
        "At every deploy: verify Cloudflare production pages with a fresh reload; a stale installed/mobile app should be guided to the single Refresh app action.",
        "Update this guide whenever routes, storage owner, database RPCs, recovery sequence or required secrets change.",
    ])
    add_callout(doc, "Final recovery reminder", "Code alone cannot recover live orders. Keep GitHub current for code and migrations, keep a recent private portable backup for data, and retain access to the Cloudflare and Supabase accounts. Together they are the complete recovery set.", PALE_ORANGE)
    # Avoid leaving the spacer paragraph after the final callout on its own page.
    final_spacer = doc.paragraphs[-1]._element
    final_spacer.getparent().remove(final_spacer)

    doc.core_properties.title = "Neeru’s Home Kitchen — Technical Recovery & Operations Guide"
    doc.core_properties.subject = "Architecture, data recovery, operations and deployment reference"
    doc.core_properties.author = "Neeru’s Home Kitchen"
    doc.core_properties.comments = "Generated from the current application source. No secrets or customer records included."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
